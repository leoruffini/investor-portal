"""
Agente Inversor PVX - Sistema automatizado de extracción de datos de inversores
y generación de Protocolos de Inversión.

Arquitectura:
  1. INPUT: Carpeta del inversor con PDFs legales (constitución, nombramiento, poderes)
  2. EXTRACCIÓN: LLM (Claude/OpenAI) analiza los PDFs y extrae datos estructurados
  3. INTERMEDIO: JSON estructurado con datos societarios, firmante, importes
  4. OUTPUT: Word del Protocolo de Inversión rellenado

Uso:
  python agente_inversor.py --carpeta "08 VENTU EUROPE (240K)" [--proyecto "Cala Corb 2"]
"""

import argparse
import json
import os
import re
import sys
from pathlib import Path
from typing import Optional

from dotenv import load_dotenv
load_dotenv(Path(__file__).resolve().parent.parent / ".env")

# PDF extraction
try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from pdf2image import convert_from_path
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

# Word generation
from docx import Document

# LLM
try:
    import anthropic
    HAS_ANTHROPIC = True
except ImportError:
    HAS_ANTHROPIC = False

try:
    import openai
    HAS_OPENAI = True
except ImportError:
    HAS_OPENAI = False


# ---------------------------------------------------------------------------
# CONFIG
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).resolve().parent.parent
TEMPLATE_PATH = BASE_DIR / "02a. Plantilla Borrador Protocolo de Inversión.docx"
OUTPUT_DIR = BASE_DIR / "output"

EXTRACTION_PROMPT = """Eres un experto en derecho mercantil español. Analiza el siguiente texto extraído de documentos legales de un inversor y extrae TODOS los datos estructurados posibles.

Devuelve un JSON con exactamente esta estructura (rellena lo que encuentres, usa null para lo que no):

{
  "datos_societarios": {
    "denominacion_actual": "nombre actual de la sociedad",
    "denominaciones_anteriores": ["lista de nombres anteriores si los hay"],
    "nif": "NIF/CIF de la sociedad",
    "forma_juridica": "S.L., S.A., etc.",
    "domicilio_social": {
      "calle": "calle y número",
      "codigo_postal": "CP",
      "localidad": "ciudad",
      "provincia": "provincia"
    },
    "objeto_social": "descripción del objeto social",
    "cnae": "código CNAE",
    "duracion": "indefinida o la que sea"
  },
  "datos_constitucion": {
    "notario": "nombre completo del notario",
    "localidad_notario": "ciudad donde se otorgó",
    "fecha": "DD de MES de AAAA",
    "numero_protocolo": "número de protocolo"
  },
  "capital_social": {
    "importe": 0.00,
    "moneda": "EUR",
    "num_participaciones": 0,
    "valor_nominal_unitario": 0.00
  },
  "datos_registrales": {
    "registro_mercantil": "localidad del RM",
    "tomo": "tomo",
    "folio": "folio",
    "hoja": "hoja",
    "inscripcion": "inscripción"
  },
  "organo_administracion": {
    "tipo": "Administrador Único / Consejo de Administración / etc.",
    "cargos": [
      {
        "cargo": "Presidente / Secretario / etc.",
        "nombre": "nombre completo",
        "documento_identidad": "DNI/NIE/Pasaporte"
      }
    ]
  },
  "representante_legal_firmante": {
    "nombre_completo": "nombre del que firma",
    "dni": "DNI/NIE",
    "cargo_en_sociedad": "cargo que ostenta",
    "domicilio": "domicilio del representante"
  }
}

IMPORTANTE:
- Extrae datos de la versión MÁS RECIENTE de cada campo (si hay cambios de nombre, usa el último)
- El representante legal firmante es quien tiene poderes para firmar en nombre de la sociedad
- Si hay conflicto entre documentos, prioriza el más reciente
- Solo devuelve el JSON, sin explicaciones adicionales

TEXTO DE LOS DOCUMENTOS:
"""


def extraer_texto_pdf(pdf_path: Path) -> str:
    """Extrae texto de un PDF. Intenta PyMuPDF primero, OCR como fallback."""
    if fitz:
        doc = fitz.open(str(pdf_path))
        text_parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        doc.close()
        full_text = "\n".join(text_parts)
        # Si hay texto sustancial, usarlo
        if len(full_text.strip()) > 200:
            return full_text

    # Fallback a OCR si PyMuPDF no extrajo texto (PDF escaneado)
    if HAS_OCR:
        from pdf2image import pdfinfo_from_path
        print(f"  [OCR] Usando OCR para {pdf_path.name}...")
        info = pdfinfo_from_path(str(pdf_path))
        num_pages = info.get("Pages", 0)
        text_parts = []
        # Process one page at a time to avoid loading all images into RAM
        for page_num in range(1, num_pages + 1):
            images = convert_from_path(
                str(pdf_path), dpi=300,
                first_page=page_num, last_page=page_num,
            )
            text = pytesseract.image_to_string(images[0], lang="spa")
            text_parts.append(text)
            del images  # free memory immediately
        return "\n".join(text_parts)

    print(f"  [WARN] No se pudo extraer texto de {pdf_path.name}. Instala PyMuPDF o pytesseract+pdf2image.")
    return ""


def buscar_pdfs(carpeta: Path) -> list[Path]:
    """Busca todos los PDFs en la carpeta del inversor."""
    pdfs = sorted(carpeta.glob("*.pdf"))
    if not pdfs:
        pdfs = sorted(carpeta.glob("**/*.pdf"))
    return pdfs


def extraer_con_llm(texto: str, llm: str = "anthropic") -> dict:
    """Envía el texto al LLM para extracción estructurada."""
    prompt = EXTRACTION_PROMPT + texto

    try:
        if llm == "anthropic" and HAS_ANTHROPIC:
            client = anthropic.Anthropic()
            response = client.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=4096,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = response.content[0].text

        elif llm == "openai" and HAS_OPENAI:
            client = openai.OpenAI()
            response = client.chat.completions.create(
                model="gpt-5.2",
                messages=[{"role": "user", "content": prompt}],
                max_completion_tokens=8192,
                reasoning_effort="medium",
            )
            raw = response.choices[0].message.content

        else:
            raise RuntimeError(
                f"LLM '{llm}' no disponible. Instala el SDK correspondiente "
                "(pip install anthropic / pip install openai)."
            )
    except Exception as e:
        if "RuntimeError" in type(e).__name__:
            raise
        raise RuntimeError(
            f"Error al llamar a la API de {llm}: {e}\n"
            "Verifica tu API key y conexión a internet."
        ) from e

    # Parsear JSON de la respuesta (buscar el primer bloque JSON balanceado)
    start = raw.find("{")
    if start == -1:
        raise ValueError(f"No se encontró JSON en la respuesta del LLM:\n{raw[:500]}")

    depth = 0
    end = start
    for i, c in enumerate(raw[start:], start):
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                end = i + 1
                break

    try:
        return json.loads(raw[start:end])
    except json.JSONDecodeError as e:
        raise ValueError(
            f"JSON inválido en respuesta del LLM (posición {e.pos}): {e.msg}\n"
            f"Fragmento: {raw[start:start+200]}..."
        ) from e


def construir_json_protocolo(datos_llm: dict, importe_inversion: float) -> dict:
    """Construye el JSON completo para el protocolo a partir de los datos del LLM."""
    ds = datos_llm.get("datos_societarios", {})
    dc = datos_llm.get("datos_constitucion", {})
    dr = datos_llm.get("datos_registrales", {})
    rl = datos_llm.get("representante_legal_firmante", {})

    # Parsear fecha de constitución
    fecha = dc.get("fecha") or ""
    dia, mes, anio = "", "", ""
    if fecha:
        parts = fecha.replace(" de ", "/").split("/")
        if len(parts) >= 3:
            dia = parts[0].strip()
            mes = parts[1].strip()
            anio = parts[2].strip()

    domicilio = ds.get("domicilio_social") or {}
    calle = domicilio.get("calle") or ""
    if domicilio.get("complemento"):
        calle += f", {domicilio['complemento']}"

    tramo_30 = importe_inversion * 0.30
    tramo_70 = importe_inversion * 0.70

    return {
        "metadata": {
            "fecha_extraccion": str(__import__("datetime").date.today()),
            "metodo": "agente_inversor_automatizado",
        },
        "datos_societarios": ds,
        "datos_constitucion": dc,
        "capital_social": datos_llm.get("capital_social", {}),
        "datos_registrales": dr,
        "organo_administracion": datos_llm.get("organo_administracion", {}),
        "representante_legal_firmante": rl,
        "campos_protocolo_inversion": {
            "inversor": {
                "sr_representante": rl.get("nombre_completo") or "",
                "denominacion_sociedad": ds.get("denominacion_actual") or "",
                "calle": calle or "",
                "numero": domicilio.get("numero") or "s/n",
                "localidad_notario_constitucion": dc.get("localidad_notario") or "",
                "nombre_notario_constitucion": dc.get("notario") or "",
                "dia_constitucion": dia or "",
                "mes_constitucion": mes or "",
                "anio_constitucion": anio or "",
                "numero_protocolo_constitucion": dc.get("numero_protocolo") or "",
                "registro_mercantil": dr.get("registro_mercantil") or "",
                "nif": ds.get("nif") or "",
            },
            "importe_inversion": {
                "total": importe_inversion,
                "total_texto": numero_a_texto(importe_inversion),
                "moneda": "EUR",
                "tramo_30_porciento": tramo_30,
                "tramo_30_texto": numero_a_texto(tramo_30),
                "tramo_70_porciento": tramo_70,
                "tramo_70_texto": numero_a_texto(tramo_70),
            },
        },
    }


def numero_a_texto(n: float) -> str:
    """Convierte un número a texto en español (simplificado)."""
    n = int(n)
    unidades = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO", "SEIS", "SIETE", "OCHO", "NUEVE"]
    decenas = ["", "DIEZ", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA",
               "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
    especiales = {
        11: "ONCE", 12: "DOCE", 13: "TRECE", 14: "CATORCE", 15: "QUINCE",
        16: "DIECISÉIS", 17: "DIECISIETE", 18: "DIECIOCHO", 19: "DIECINUEVE",
        21: "VEINTIÚN", 22: "VEINTIDÓS", 23: "VEINTITRÉS", 24: "VEINTICUATRO",
        25: "VEINTICINCO", 26: "VEINTISÉIS", 27: "VEINTISIETE", 28: "VEINTIOCHO",
        29: "VEINTINUEVE",
    }
    centenas = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
                "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]

    if n == 0:
        return "CERO"
    if n == 100:
        return "CIEN"

    parts = []

    if n >= 1_000_000:
        millones = n // 1_000_000
        if millones == 1:
            parts.append("UN MILLÓN")
        else:
            parts.append(f"{numero_a_texto(millones)} MILLONES")
        n %= 1_000_000

    if n >= 1000:
        miles = n // 1000
        if miles == 1:
            parts.append("MIL")
        else:
            parts.append(f"{numero_a_texto(miles)} MIL")
        n %= 1000

    if n >= 100:
        if n == 100:
            parts.append("CIEN")
            return " ".join(parts)
        parts.append(centenas[n // 100])
        n %= 100

    if n > 0:
        if n in especiales:
            parts.append(especiales[n])
        elif n < 10:
            parts.append(unidades[n])
        else:
            d = decenas[n // 10]
            u = unidades[n % 10]
            if u:
                parts.append(f"{d} Y {u}")
            else:
                parts.append(d)

    return " ".join(parts)


def rellenar_word(template_path: Path, datos: dict, output_path: Path):
    """Rellena la plantilla Word con los datos del JSON (misma lógica del PoC)."""
    doc = Document(str(template_path))
    inv = datos["campos_protocolo_inversion"]["inversor"]
    imp = datos["campos_protocolo_inversion"]["importe_inversion"]

    def reemplazar_secuencial(paragraph, valores):
        full_text = paragraph.text
        for valor in valores:
            full_text = full_text.replace("[*]", str(valor), 1)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    inversor_parrafo_values = [
        inv["sr_representante"], inv["denominacion_sociedad"],
        inv["calle"], inv["numero"],
        inv["localidad_notario_constitucion"], inv["nombre_notario_constitucion"],
        inv["dia_constitucion"], inv["mes_constitucion"], inv["anio_constitucion"],
        inv["numero_protocolo_constitucion"],
    ]
    inversor_registro_values = [inv["registro_mercantil"], inv["nif"]]
    inversor_asegura_values = [inv["sr_representante"]]

    # Rellenar párrafos del inversor
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '"el Inversor"' in text and "en nombre y representación" in text:
            reemplazar_secuencial(para, inversor_parrafo_values)
        elif text.strip().startswith("Sr. [*]") and i > 50:
            full = para.text
            first_pos = full.find("[*]")
            if first_pos >= 0:
                second_pos = full.find("[*]", first_pos + 3)
                if second_pos >= 0:
                    new_text = full[:second_pos] + (inv["sr_representante"] or "") + full[second_pos + 3:]
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ""
        elif "PROVALIX PROMOCIONES" in text and i > 55 and text.count("[*]") >= 2:
            full = para.text
            last_pos = full.rfind("[*]")
            if last_pos >= 0:
                new_text = full[:last_pos] + (inv["denominacion_sociedad"] or "") + full[last_pos + 3:]
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""

    # Registro e inscripción del inversor (segundo de cada tipo)
    registro_count = 0
    asegura_count = 0
    for para in doc.paragraphs:
        text = para.text
        if "Inscrita en el Registro Mercantil" in text and "NIF" in text:
            registro_count += 1
            if registro_count == 2:
                reemplazar_secuencial(para, inversor_registro_values)
        if "asegura que el poder y facultades" in text:
            asegura_count += 1
            if asegura_count == 2:
                reemplazar_secuencial(para, inversor_asegura_values)

    # Importes de inversión en párrafos
    for para in doc.paragraphs:
        text = para.text
        if "la sociedad" in text and "el Inversor" in text and "aportación de" in text:
            reemplazar_secuencial(para, [
                inv["denominacion_sociedad"], "[Nombre del Proyecto]",
                imp["total_texto"],
                f"{imp['total']:,.0f}".replace(",", "."),
            ])
        if "equivalente al 30%" in text:
            reemplazar_secuencial(para, [
                imp["tramo_30_texto"],
                f"{imp['tramo_30_porciento']:,.0f}".replace(",", "."),
            ])
        if "equivalente al 70%" in text and "segunda Nota Convertible" in text:
            reemplazar_secuencial(para, [
                imp["tramo_70_texto"],
                f"{imp['tramo_70_porciento']:,.0f}".replace(",", "."),
            ])

    # --- RELLENAR TABLAS CON IMPORTES ---
    def reemplazar_importe_en_parrafo(paragraph, texto_importe, cifra_importe):
        """Reemplaza '[*] EUROS ([*] €)' sin tocar otros [*] (ej: PROVALIX [*])."""
        text = paragraph.text
        if "[*] EUROS" not in text:
            return
        new_text = text.replace("[*] EUROS", f"{texto_importe} EUROS", 1)
        new_text = new_text.replace("([*] €)", f"({cifra_importe} €)", 1)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_lower = cell.text.lower()

                # C.2 y C.3: importes 30%
                if ("prestamista" in cell_lower and "importe de" in cell_lower) or \
                   "desembolsa en este acto" in cell_lower:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph, imp["tramo_30_texto"],
                            f"{imp['tramo_30_porciento']:,.0f}".replace(",", "."),
                        )

                # D.1 y D.2: importes 70%
                if ("segunda aportación" in cell_lower or "incumplimiento" in cell_lower) \
                   and "70%" in cell.text:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph, imp["tramo_70_texto"],
                            f"{imp['tramo_70_porciento']:,.0f}".replace(",", "."),
                        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def extraer_importe_de_carpeta(nombre_carpeta: str) -> Optional[float]:
    """Intenta extraer el importe de inversión del nombre de la carpeta (ej: '08 VENTU EUROPE (240K)')."""
    match = re.search(r"\((\d+)[Kk]\)", nombre_carpeta)
    if match:
        return float(match.group(1)) * 1000
    match = re.search(r"\((\d[\d.]+)\)", nombre_carpeta)
    if match:
        return float(match.group(1).replace(".", ""))
    return None


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------
def main():
    parser = argparse.ArgumentParser(
        description="Agente Inversor PVX: extrae datos de PDFs legales y genera Protocolo de Inversión"
    )
    parser.add_argument(
        "--carpeta", required=True,
        help="Nombre de la carpeta del inversor (ej: '08 VENTU EUROPE (240K)')"
    )
    parser.add_argument(
        "--importe", type=float, default=None,
        help="Importe de inversión en EUR. Si no se indica, se intenta extraer del nombre de carpeta."
    )
    parser.add_argument(
        "--llm", default="anthropic", choices=["anthropic", "openai"],
        help="LLM a usar para la extracción (default: anthropic)"
    )
    parser.add_argument(
        "--template", default=str(TEMPLATE_PATH),
        help="Ruta a la plantilla Word del protocolo"
    )
    parser.add_argument(
        "--output-dir", default=str(OUTPUT_DIR),
        help="Directorio de salida"
    )
    args = parser.parse_args()

    carpeta = BASE_DIR / args.carpeta
    if not carpeta.exists():
        print(f"ERROR: Carpeta no encontrada: {carpeta}")
        sys.exit(1)

    # Importe
    importe = args.importe or extraer_importe_de_carpeta(args.carpeta)
    if not importe:
        print("ERROR: No se pudo determinar el importe de inversión. Usa --importe.")
        sys.exit(1)
    print(f"Importe de inversión: {importe:,.0f} EUR")

    # Nombre del inversor (de la carpeta)
    nombre_inversor = re.sub(r"^\d+\s*", "", args.carpeta)
    nombre_inversor = re.sub(r"\s*\(.*\)", "", nombre_inversor).strip()
    print(f"Inversor: {nombre_inversor}")

    # 1. Buscar PDFs
    pdfs = buscar_pdfs(carpeta)
    if not pdfs:
        print(f"ERROR: No se encontraron PDFs en {carpeta}")
        sys.exit(1)
    print(f"\nPDFs encontrados ({len(pdfs)}):")
    for p in pdfs:
        print(f"  - {p.name}")

    # 2. Extraer texto
    print("\nExtrayendo texto de PDFs...")
    textos = []
    for pdf in pdfs:
        print(f"  Procesando: {pdf.name}...")
        texto = extraer_texto_pdf(pdf)
        if texto:
            textos.append(f"=== DOCUMENTO: {pdf.name} ===\n{texto}")
            print(f"    -> {len(texto):,} caracteres extraídos")
        else:
            print(f"    -> Sin texto extraíble")

    texto_completo = "\n\n".join(textos)

    # Limitar texto si es muy largo (para no exceder contexto del LLM)
    MAX_CHARS = 400_000
    if len(texto_completo) > MAX_CHARS:
        print(f"\n[WARN] Texto total ({len(texto_completo):,} chars) excede límite. Truncando a {MAX_CHARS:,}.")
        texto_completo = texto_completo[:MAX_CHARS]

    # 3. Extraer datos con LLM
    print(f"\nExtrayendo datos con {args.llm}...")
    datos_llm = extraer_con_llm(texto_completo, args.llm)

    # 4. Construir JSON del protocolo
    datos_protocolo = construir_json_protocolo(datos_llm, importe)

    # 5. Guardar JSON
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    nombre_safe = nombre_inversor.replace(" ", "_").upper()
    json_path = output_dir / f"{nombre_safe}_datos_extraidos.json"
    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(datos_protocolo, f, ensure_ascii=False, indent=2)
    print(f"\nJSON guardado: {json_path}")

    # 6. Generar Word
    template = Path(args.template)
    if not template.exists():
        print(f"WARN: Plantilla no encontrada: {template}. Solo se genera JSON.")
    else:
        word_path = output_dir / f"Protocolo_Inversion_{nombre_safe}.docx"
        rellenar_word(template, datos_protocolo, word_path)
        print(f"Protocolo Word generado: {word_path}")

    print("\nProceso completado.")


if __name__ == "__main__":
    main()
