"""
Script para rellenar la plantilla del Protocolo de Inversión con datos del inversor.
PoC: VENTU EUROPE (240K)
"""

import json
from pathlib import Path
from docx import Document


def cargar_datos(json_path: str) -> dict:
    with open(json_path, "r", encoding="utf-8") as f:
        return json.load(f)


def reemplazar_en_run(run, placeholder: str, valor: str):
    """Reemplaza un placeholder en un run individual."""
    if placeholder in run.text:
        run.text = run.text.replace(placeholder, valor)


def reemplazar_en_parrafo(paragraph, placeholder: str, valor: str):
    """Reemplaza placeholder en un párrafo, manejando runs fragmentados."""
    full_text = paragraph.text
    if placeholder not in full_text:
        return False

    # Intentar reemplazo directo en cada run
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, valor)
            return True

    # Si el placeholder está fragmentado entre runs, reconstruir
    new_text = full_text.replace(placeholder, valor)
    # Limpiar todos los runs y poner el texto en el primero
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def reemplazar_en_tabla(table, placeholder: str, valor: str):
    """Reemplaza placeholder en todas las celdas de una tabla."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                reemplazar_en_parrafo(paragraph, placeholder, valor)


def reemplazar_global(doc: Document, placeholder: str, valor: str):
    """Reemplaza un placeholder en todo el documento (párrafos, tablas, headers, footers)."""
    for paragraph in doc.paragraphs:
        reemplazar_en_parrafo(paragraph, placeholder, valor)

    for table in doc.tables:
        reemplazar_en_tabla(table, placeholder, valor)

    for section in doc.sections:
        if section.header:
            for paragraph in section.header.paragraphs:
                reemplazar_en_parrafo(paragraph, placeholder, valor)
        if section.footer:
            for paragraph in section.footer.paragraphs:
                reemplazar_en_parrafo(paragraph, placeholder, valor)


def contar_placeholders(doc: Document) -> list:
    """Cuenta los [*] restantes en el documento."""
    restantes = []
    for i, para in enumerate(doc.paragraphs):
        count = para.text.count("[*]")
        if count > 0:
            restantes.append(f"  Párrafo {i}: {count}x [*] -> {para.text[:100]}...")
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for paragraph in cell.paragraphs:
                    count = paragraph.text.count("[*]")
                    if count > 0:
                        restantes.append(
                            f"  Tabla {t_idx}, Fila {r_idx}, Col {c_idx}: {count}x [*] -> {paragraph.text[:100]}..."
                        )
    return restantes


def rellenar_protocolo(template_path: str, json_path: str, output_path: str):
    """Rellena la plantilla del protocolo con los datos del JSON."""
    datos = cargar_datos(json_path)
    doc = Document(template_path)

    inv = datos["campos_protocolo_inversion"]["inversor"]
    imp = datos["campos_protocolo_inversion"]["importe_inversion"]

    # --- RELLENAR CAMPOS DEL INVERSOR ---
    # El template tiene párrafos 27-28 para el inversor con múltiples [*].
    # Necesitamos reemplazar los [*] en orden dentro de esos párrafos específicos.

    # Párrafo 27: "El Sr. [*] en nombre y representación de la sociedad [*]
    #   (en adelante, también, "el Inversor"), domiciliada en la calle [*], número [*],
    #   constituida por tiempo indefinido mediante escritura autorizada por el Notario de [*],
    #   Don [*], el día [*] de [*] de [*], con el número [*] de su protocolo."
    inversor_parrafo_values = [
        inv["sr_representante"],                     # Sr. [*]
        inv["denominacion_sociedad"],                # sociedad [*]
        inv["calle"],                                # calle [*]
        inv["numero"],                               # número [*]
        inv["localidad_notario_constitucion"],        # Notario de [*]
        inv["nombre_notario_constitucion"],           # Don [*]
        inv["dia_constitucion"],                      # día [*]
        inv["mes_constitucion"],                      # de [*]
        inv["anio_constitucion"],                     # de [*]
        inv["numero_protocolo_constitucion"],         # número [*]
    ]

    # Párrafo 28: "Inscrita en el Registro Mercantil de [*], y provista de NIF. número [*]."
    inversor_registro_values = [
        inv["registro_mercantil"],                   # Registro Mercantil de [*]
        inv["nif"],                                  # NIF. número [*]
    ]

    # Párrafo 29: "El Sr. [*] asegura que..."
    inversor_asegura_values = [
        inv["sr_representante"],                     # Sr. [*]
    ]

    # Reemplazar [*] secuencialmente en párrafos específicos del inversor
    def reemplazar_secuencial(paragraph, valores):
        """Reemplaza [*] en orden dentro de un párrafo."""
        full_text = paragraph.text
        for valor in valores:
            full_text = full_text.replace("[*]", valor, 1)
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""

    # Identificar los párrafos del inversor (27, 28, 29 según la extracción)
    for i, para in enumerate(doc.paragraphs):
        text = para.text

        # Párrafo del inversor (contiene "el Inversor")
        if '"el Inversor"' in text and "en nombre y representación" in text:
            reemplazar_secuencial(para, inversor_parrafo_values)

        # Registro del inversor (segundo "Inscrita en el Registro")
        # Hay dos: uno para Provalix (párrafo 25) y otro para el Inversor (párrafo 28)
        # El del inversor viene justo después del párrafo del inversor
        elif "Inscrita en el Registro Mercantil" in text and "NIF. número" in text:
            # Marcamos el segundo como inversor (el primero es Provalix)
            pass  # Se maneja abajo

        # Párrafo "asegura" del inversor
        elif "asegura que el poder y facultades" in text:
            pass  # Se maneja abajo

        # Párrafo firma: "Sr. [*]		     		                               Sr. [*]"
        # El primer [*] es Provalix (dejar), el segundo es Inversor
        elif text.strip().startswith("Sr. [*]") and i > 50:
            # Solo reemplazar el SEGUNDO [*] (inversor), dejar el primero (Provalix)
            full = para.text
            first_pos = full.find("[*]")
            if first_pos >= 0:
                second_pos = full.find("[*]", first_pos + 3)
                if second_pos >= 0:
                    # Reemplazar solo el segundo
                    new_text = full[:second_pos] + inv["sr_representante"] + full[second_pos + 3:]
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ""

        # Párrafo firma denominación: "PROVALIX PROMOCIONES [*], S.L.		[*]"
        # El primer [*] es Provalix (dejar), el segundo es Inversor
        elif "PROVALIX PROMOCIONES" in text and i > 55 and text.count("[*]") >= 2:
            full = para.text
            # Reemplazar solo el último [*]
            last_pos = full.rfind("[*]")
            if last_pos >= 0:
                new_text = full[:last_pos] + inv["denominacion_sociedad"] + full[last_pos + 3:]
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""

        # Párrafo fecha: "En Barcelona, a [*] de [*] de [*]"
        # Estos son fecha del acto - NO rellenar con datos del inversor (dejar como [*])

    # Ahora manejo los párrafos de Registro y Asegura:
    # Necesito distinguir entre los de Provalix (primeros) y los del Inversor (segundos)
    registro_count = 0
    asegura_count = 0
    for para in doc.paragraphs:
        text = para.text
        if "Inscrita en el Registro Mercantil" in text and "NIF" in text:
            registro_count += 1
            if registro_count == 2:  # Segundo = Inversor
                reemplazar_secuencial(para, inversor_registro_values)
        if "asegura que el poder y facultades" in text:
            asegura_count += 1
            if asegura_count == 2:  # Segundo = Inversor
                reemplazar_secuencial(para, inversor_asegura_values)

    # --- RELLENAR CAMPOS DE IMPORTE DE INVERSIÓN ---
    # En tablas y párrafos del Objeto
    for para in doc.paragraphs:
        text = para.text

        # Párrafo 33: inversión del inversor en proyecto
        if "la sociedad" in text and "el Inversor" in text and "Inversión" in text and "aportación de" in text:
            reemplazar_secuencial(para, [
                inv["denominacion_sociedad"],           # sociedad [*]
                "[Nombre del Proyecto]",                # Proyecto de [*]
                f"{imp['total_texto']}",                # [*] EUROS
                f"{imp['total']:,.0f}".replace(",", "."),  # ([*] €)
            ])

        # Párrafo 35: 30% de la inversión
        if "equivalente al 30%" in text:
            reemplazar_secuencial(para, [
                f"{imp['tramo_30_texto']}",
                f"{imp['tramo_30_porciento']:,.0f}".replace(",", "."),
            ])

        # Párrafo 36: 70% de la inversión
        if "equivalente al 70%" in text and "segunda Nota Convertible" in text:
            reemplazar_secuencial(para, [
                f"{imp['tramo_70_texto']}",
                f"{imp['tramo_70_porciento']:,.0f}".replace(",", "."),
            ])

    # --- RELLENAR TABLAS CON DATOS DEL INVERSOR ---
    # Estrategia: usar reemplazo selectivo por patrón "[*] EUROS" y "([*] €)"
    # en vez de secuencial, para evitar pisar el [*] de PROVALIX PROMOCIONES [*]
    def reemplazar_importe_en_parrafo(paragraph, texto_importe, cifra_importe):
        """Reemplaza '[*] EUROS ([*] €)' con los valores correctos sin tocar otros [*]."""
        text = paragraph.text
        if "[*] EUROS" not in text:
            return
        # Reemplazar el patrón "[*] EUROS" -> "TEXTO EUROS"
        new_text = text.replace("[*] EUROS", f"{texto_importe} EUROS", 1)
        # Reemplazar el patrón "([*] €)" -> "(CIFRA €)"
        new_text = new_text.replace("([*] €)", f"({cifra_importe} €)", 1)
        if paragraph.runs:
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text_lower = cell.text.lower()

                # Tabla C.2: Consentimiento - importe del préstamo convertible (30%)
                if "prestamista" in cell_text_lower and "importe de" in cell_text_lower:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph,
                            imp['tramo_30_texto'],
                            f"{imp['tramo_30_porciento']:,.0f}".replace(",", "."),
                        )

                # Tabla C.3: Desembolso 30%
                if "desembolsa en este acto" in cell_text_lower:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph,
                            imp['tramo_30_texto'],
                            f"{imp['tramo_30_porciento']:,.0f}".replace(",", "."),
                        )

                # Tabla D.1: Segunda aportación 70%
                if "segunda aportación" in cell_text_lower and "70%" in cell.text:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph,
                            imp['tramo_70_texto'],
                            f"{imp['tramo_70_porciento']:,.0f}".replace(",", "."),
                        )

                # Tabla D.2: Consecuencias incumplimiento
                if "incumplimiento" in cell_text_lower and "70%" in cell.text:
                    for paragraph in cell.paragraphs:
                        reemplazar_importe_en_parrafo(
                            paragraph,
                            imp['tramo_70_texto'],
                            f"{imp['tramo_70_porciento']:,.0f}".replace(",", "."),
                        )

    # --- GUARDAR ---
    doc.save(output_path)

    # Reportar placeholders restantes
    restantes = contar_placeholders(doc)
    print(f"Protocolo generado: {output_path}")
    print(f"Placeholders [*] rellenados del inversor.")
    if restantes:
        print(f"\n[*] restantes ({len(restantes)} ubicaciones) - corresponden a datos de Provalix/Proyecto:")
        for r in restantes:
            print(r)
    else:
        print("Todos los placeholders han sido rellenados.")


if __name__ == "__main__":
    base = Path("/Users/leoruffini/Desktop/PVX - Agente Inversor")
    template = base / "02a. Plantilla Borrador Protocolo de Inversión.docx"
    datos_json = base / "output" / "ventu_europe_datos_extraidos.json"
    output = base / "output" / "Protocolo_Inversion_VENTU_EUROPE.docx"

    output.parent.mkdir(parents=True, exist_ok=True)
    rellenar_protocolo(str(template), str(datos_json), str(output))
