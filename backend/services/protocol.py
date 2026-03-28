"""
Protocol generation service.

Wraps the logic from scripts/rellenar_protocolo.py:
  KYC JSON + Word template → filled Investment Protocol (.docx)
"""

import io
from pathlib import Path

from docx import Document


BASE_DIR = Path(__file__).resolve().parent.parent.parent
TEMPLATE_PATH = BASE_DIR / "02a. Plantilla Borrador Protocolo de Inversión.docx"


def _reemplazar_en_parrafo(paragraph, placeholder: str, valor: str) -> bool:
    """Replace placeholder in a paragraph, handling fragmented runs."""
    full_text = paragraph.text
    if placeholder not in full_text:
        return False

    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, valor)
            return True

    # Placeholder is fragmented across runs — reconstruct
    new_text = full_text.replace(placeholder, valor)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    return True


def _reemplazar_secuencial(paragraph, valores: list[str]):
    """Replace [*] placeholders sequentially in a paragraph."""
    full_text = paragraph.text
    for valor in valores:
        full_text = full_text.replace("[*]", str(valor), 1)
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _reemplazar_importe_en_parrafo(paragraph, texto_importe: str, cifra_importe: str):
    """Replace '[*] EUROS ([*] €)' without touching other [*] placeholders."""
    text = paragraph.text
    if "[*] EUROS" not in text:
        return
    new_text = text.replace("[*] EUROS", f"{texto_importe} EUROS", 1)
    new_text = new_text.replace("([*] €)", f"({cifra_importe} €)", 1)
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""


def generate_protocol(datos: dict, template_path: Path | None = None) -> bytes:
    """
    Fill the Word template with protocol data and return the document as bytes.

    Args:
        datos: Full protocol JSON (output of build_protocol_json).
        template_path: Path to .docx template. Defaults to repo template.

    Returns:
        The filled .docx file as bytes.
    """
    template = template_path or TEMPLATE_PATH
    if not template.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {template}")

    doc = Document(str(template))

    inv = datos["campos_protocolo_inversion"]["inversor"]
    imp = datos["campos_protocolo_inversion"]["importe_inversion"]

    # Format amounts for display
    primer_cifra = f"{imp['primer_desembolso']:,.0f}".replace(",", ".")
    segundo_cifra = f"{imp['segundo_desembolso']:,.0f}".replace(",", ".")
    total_cifra = f"{imp['total']:,.0f}".replace(",", ".")
    primer_pct = int(imp["primer_desembolso_pct"]) if imp["primer_desembolso_pct"] == int(imp["primer_desembolso_pct"]) else imp["primer_desembolso_pct"]
    segundo_pct = int(imp["segundo_desembolso_pct"]) if imp["segundo_desembolso_pct"] == int(imp["segundo_desembolso_pct"]) else imp["segundo_desembolso_pct"]

    # --- Investor paragraph values ---
    inversor_parrafo_values = [
        inv["sr_representante"], inv["denominacion_sociedad"],
        inv["calle"], inv["numero"],
        inv["localidad_notario_constitucion"], inv["nombre_notario_constitucion"],
        inv["dia_constitucion"], inv["mes_constitucion"], inv["anio_constitucion"],
        inv["numero_protocolo_constitucion"],
    ]
    inversor_registro_values = [inv["registro_mercantil"], inv["nif"]]
    inversor_asegura_values = [inv["sr_representante"]]

    # --- Fill investor paragraphs ---
    for i, para in enumerate(doc.paragraphs):
        text = para.text

        if '"el Inversor"' in text and "en nombre y representación" in text:
            _reemplazar_secuencial(para, inversor_parrafo_values)

        elif text.strip().startswith("Sr. [*]") and i > 50:
            full = para.text
            first_pos = full.find("[*]")
            if first_pos >= 0:
                second_pos = full.find("[*]", first_pos + 3)
                if second_pos >= 0:
                    new_text = full[:second_pos] + (inv["sr_representante"] or "") + full[second_pos + 3:]
                    if para.runs:
                        para.runs[0].text = new_text
                        for run in para.runs[1:]:
                            run.text = ""

        elif "PROVALIX PROMOCIONES" in text and i > 55 and text.count("[*]") >= 2:
            full = para.text
            last_pos = full.rfind("[*]")
            if last_pos >= 0:
                new_text = full[:last_pos] + (inv["denominacion_sociedad"] or "") + full[last_pos + 3:]
                if para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""

    # Registry and "asegura" paragraphs (second occurrence = investor)
    registro_count = 0
    asegura_count = 0
    for para in doc.paragraphs:
        text = para.text
        if "Inscrita en el Registro Mercantil" in text and "NIF" in text:
            registro_count += 1
            if registro_count == 2:
                _reemplazar_secuencial(para, inversor_registro_values)
        if "asegura que el poder y facultades" in text:
            asegura_count += 1
            if asegura_count == 2:
                _reemplazar_secuencial(para, inversor_asegura_values)

    # --- Fill investment amounts in paragraphs ---
    # Build dynamic percentage strings for matching (e.g., "30%", "50%")
    first_pct_str = f"{primer_pct}%"
    second_pct_str = f"{segundo_pct}%"

    for para in doc.paragraphs:
        text = para.text

        if "la sociedad" in text and "el Inversor" in text and "aportación de" in text:
            _reemplazar_secuencial(para, [
                inv["denominacion_sociedad"], "[Nombre del Proyecto]",
                imp["total_texto"], total_cifra,
            ])

        if f"equivalente al {first_pct_str}" in text and "segunda" not in text.lower():
            _reemplazar_secuencial(para, [
                imp["primer_desembolso_texto"], primer_cifra,
            ])

        if f"equivalente al {second_pct_str}" in text and "segunda" in text.lower():
            _reemplazar_secuencial(para, [
                imp["segundo_desembolso_texto"], segundo_cifra,
            ])

    # --- Fill investment amounts in tables ---
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_lower = cell.text.lower()

                if ("prestamista" in cell_lower and "importe de" in cell_lower) or \
                   "desembolsa en este acto" in cell_lower:
                    for paragraph in cell.paragraphs:
                        _reemplazar_importe_en_parrafo(
                            paragraph, imp["primer_desembolso_texto"], primer_cifra,
                        )

                if ("segunda aportación" in cell_lower or "incumplimiento" in cell_lower) \
                   and second_pct_str in cell.text:
                    for paragraph in cell.paragraphs:
                        _reemplazar_importe_en_parrafo(
                            paragraph, imp["segundo_desembolso_texto"], segundo_cifra,
                        )

    # Return as bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
